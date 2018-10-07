# coding=utf-8

import json
import os
import docx
from docx.oxml.ns import qn

DIR_TO_IMPORT = 'notes'

def save_file(file_name, date, content):
    # file name
    file_name = os.path.join(DIR_TO_IMPORT, file_name) + '.docx'

    # content
    content = '\n'.join([d.strip() for d in content.split('\n')[1:]])

    # save as word
    document = docx.Document()
    document.styles['Normal'].font.name = u'宋体'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    document.add_paragraph(content)
    document.add_paragraph(date)
    document.save(file_name)


def mkdir(path, prefix=DIR_TO_IMPORT):
    if not os.path.exists(os.path.join(prefix, path)):
        os.makedirs(os.path.join(prefix, path))


def parse_items(items, box):
    for d in items:
        date = d['created_time']
        date = date.replace(':', '点')
        date += "分"
        content = d['content']
        save_file(os.path.join(box, date), date, content)

def run():
    mkdir(DIR_TO_IMPORT, '')
    with open('data.json', mode='r', encoding='utf8') as fp:
        data = json.loads(fp.read())
        singles = []
        for item in data:
            if item.get('box_name'):  # if classification
                mkdir(item['box_name'])
                parse_items(item['content'],item['box_name'])
            else:
                singles.append(item)
        parse_items(singles, '')

if __name__ == '__main__':
    run()
